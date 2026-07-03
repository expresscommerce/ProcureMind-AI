import io
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

def extract_text_and_tables(file_bytes: bytes, file_type: str) -> str:
    extracted_text = ""
    
    if file_type == "application/pdf" or file_type.endswith("pdf"):
        # PyMuPDF for text
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            extracted_text += f"\n--- Page {page_num + 1} ---\n"
            extracted_text += page.get_text("text")
        
        # pdfplumber for tables
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    extracted_text += f"\n--- Tables on Page {i + 1} ---\n"
                    for table in tables:
                        for row in table:
                            extracted_text += " | ".join([str(cell).strip() if cell else "" for cell in row]) + "\n"
                        extracted_text += "\n"
                        
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_type.endswith("docx"):
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            extracted_text += para.text + "\n"
            
        for table in doc.tables:
            extracted_text += "\n--- Table ---\n"
            for row in table.rows:
                extracted_text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
            extracted_text += "\n"
            
    else:
        # Fallback for plain text or unknown
        try:
            extracted_text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            extracted_text = ""

    return extracted_text
